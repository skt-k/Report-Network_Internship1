import pandas as pd
from openpyxl.chart import BarChart, LineChart
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ALL_COLS = [
    "room_point",
    "note",
    "floor",
    "ssid",
    "bssid",
    "ap_vendor",
    "band",
    "radio_type",
    "channel",
    "signal_percent",
    "rssi_dbm",
    "receive_rate_mbps",
    "transmit_rate_mbps",
    "gateway_ip",
    "ping_gateway_ms",
    "ping_gateway_loss_pct",
    "server_ip",
    "trace_target",
    "ping_server_ms",
    "ping_server_loss_pct",
    "tcp_upload_mbps",
    "tcp_download_mbps",
    "udp_target_bandwidth",
    "udp_actual_mbps",
    "udp_jitter_ms",
    "udp_packetloss_pct",
    "co_channel_ap_count",
    "adjacent_ap_count",
    "strongest_neighbor_rssi",
    "noise_floor_dbm",
    "snr_db",
    "snr_quality",
    "rating",
]

ALL_COLS_TH = [
    "ห้อง",
    "จุด",
    "ชั้น",
    "SSID",
    "BSSID",
    "AP Vendor",
    "Band",
    "Radio Type",
    "Channel",
    "Signal %",
    "RSSI (dBm)",
    "Receive Rate (Mbps)",
    "Transmit Rate (Mbps)",
    "Gateway IP",
    "Ping Gateway (ms)",
    "Gateway Loss %",
    "Server IP",
    "Trace Target",
    "Ping Server (ms)",
    "Server Loss %",
    "TCP Upload (Mbps)",
    "TCP Download (Mbps)",
    "UDP Target BW",
    "UDP Actual (Mbps)",
    "UDP Jitter (ms)",
    "UDP Loss %",
    "Co-Channel AP",
    "Adjacent AP",
    "Neighbor RSSI",
    "Noise Floor (dBm)",
    "SNR (dB)",
    "SNR Quality",
    "Rating",
]

FIELD_MAP = [
    ("SSID", "ssid"),
    ("BSSID", "bssid"),
    ("AP Vendor", "ap_vendor"),
    ("Band", "band"),
    ("Radio Type", "radio_type"),
    ("Channel", "channel"),
    ("Signal %", "signal_percent"),
    ("RSSI (dBm)", "rssi_dbm"),
    ("Receive Rate (Mbps)", "receive_rate_mbps"),
    ("Transmit Rate (Mbps)", "transmit_rate_mbps"),
    ("Gateway IP", "gateway_ip"),
    ("Ping Gateway (ms)", "ping_gateway_ms"),
    ("Gateway Loss %", "ping_gateway_loss_pct"),
    ("Server IP", "server_ip"),
    ("Trace Target", "trace_target"),
    ("Ping Server (ms)", "ping_server_ms"),
    ("Server Loss %", "ping_server_loss_pct"),
    ("TCP Upload (Mbps)", "tcp_upload_mbps"),
    ("TCP Download (Mbps)", "tcp_download_mbps"),
    ("UDP Target BW (Mbps)", "udp_target_bandwidth"),
    ("UDP Actual (Mbps)", "udp_actual_mbps"),
    ("UDP Jitter (ms)", "udp_jitter_ms"),
    ("UDP Packet Loss %", "udp_packetloss_pct"),
    ("Co-Channel AP", "co_channel_ap_count"),
    ("Adjacent AP", "adjacent_ap_count"),
    ("Strongest Neighbor RSSI", "strongest_neighbor_rssi"),
    ("Rating", "rating"),
]

MANUAL_BLANK_COLS = {
    "co_channel_ap_count",
    "strongest_neighbor_rssi",
    "adjacent_ap_count",
}

HEADER_FILL = PatternFill("solid", fgColor="2C3E50")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=9, name="TH Sarabun New")
DATA_FONT = Font(size=9, name="TH Sarabun New")

RATING_FILLS = {
    "Good": PatternFill("solid", fgColor="D5F5E3"),
    "Excellent": PatternFill("solid", fgColor="A9DFBF"),
    "Poor": PatternFill("solid", fgColor="FADBD8"),
}


def safe_float(value):
    try:
        if value is None or (isinstance(value, float) and pd.isna(value)):
            return None
        return float(value)
    except Exception:
        return None


def safe_val(value, decimals=None):
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "-"
    if decimals is not None:
        try:
            return str(round(float(value), decimals))
        except Exception:
            return str(value)
    return str(value)


def thin_border():
    side = Side(style="thin", color="CCCCCC")
    return Border(left=side, right=side, top=side, bottom=side)


def shade_cell(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def style_word_header(cell, text, bg="1B2A3B", fg="FFFFFF"):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.italic = False
    run.font.underline = False
    run.font.color.rgb = RGBColor.from_string(fg)
    run.font.size = Pt(10)
    run.font.name = "TH Sarabun New"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, bg)


def make_bar_chart(
    title,
    y_title,
    x_title,
    y_min=None,
    y_max=None,
    major_unit=None,
    width=26,
    height=14,
):
    chart = BarChart()
    chart.type = "col"
    chart.grouping = "clustered"
    chart.title = title
    chart.style = 2
    chart.width = width
    chart.height = height
    chart.y_axis.title = y_title
    chart.y_axis.numFmt = "0"
    chart.y_axis.delete = False
    chart.y_axis.majorGridlines = None
    chart.y_axis.tickLblPos = "low"
    if y_min is not None:
        chart.y_axis.scaling.min = y_min
    if y_max is not None:
        chart.y_axis.scaling.max = y_max
    if major_unit is not None:
        chart.y_axis.majorUnit = major_unit
    chart.x_axis.title = x_title
    chart.x_axis.tickLblPos = "low"
    chart.x_axis.noMultiLvlLbl = True
    return chart


def make_line_chart(title, y_title, x_title, y_min=0, width=26, height=14):
    chart = LineChart()
    chart.title = title
    chart.style = 2
    chart.width = width
    chart.height = height
    chart.y_axis.title = y_title
    chart.y_axis.numFmt = "0.0"
    chart.y_axis.delete = False
    chart.y_axis.majorGridlines = None
    chart.y_axis.tickLblPos = "low"
    if y_min is not None:
        chart.y_axis.scaling.min = y_min
    chart.x_axis.title = x_title
    chart.x_axis.tickLblPos = "low"
    chart.x_axis.noMultiLvlLbl = True
    return chart
