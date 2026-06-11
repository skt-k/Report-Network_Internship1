from dotenv import load_dotenv
import os
import pandas as pd
import numpy as np
import openpyxl
import tkinter as tk
from tkinter import ttk, messagebox
from supabase import create_client
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

# ============================================================
# ดึงรายชื่อตึกจาก Supabase แล้วแสดง Popup ให้เลือก
# ============================================================
client = create_client(SUPABASE_URL, SUPABASE_KEY)

buildings_resp = client.table("surveys").select("building").execute()
buildings = sorted(set(row["building"] for row in buildings_resp.data if row.get("building")))

if not buildings:
    print("❌ ไม่พบข้อมูลตึกใน Supabase")
    exit()

def select_building(buildings):
    result = []

    root = tk.Tk()
    root.title("เลือกตึก")
    root.geometry("320x140")
    root.resizable(False, False)
    root.eval("tk::PlaceWindow . center")

    tk.Label(root, text="เลือกตึกที่ต้องการ generate report:", font=("TH Sarabun New", 11)).pack(pady=12)

    selected = tk.StringVar(value=buildings[0])
    dropdown = ttk.Combobox(root, textvariable=selected, values=buildings, state="readonly", width=28, font=("TH Sarabun New", 11))
    dropdown.pack()

    def confirm():
        result.append(selected.get())
        root.destroy()

    def on_close():
        root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_close)
    tk.Button(root, text="ตกลง", command=confirm, width=12, font=("TH Sarabun New", 10)).pack(pady=12)
    root.mainloop()

    return result[0] if result else None

TARGET_BUILDING = select_building(buildings)

if not TARGET_BUILDING:
    print("❌ ไม่ได้เลือกตึก หยุดการทำงาน")
    exit()

# ============================================================
# ดึงข้อมูล
# ============================================================
response = client.table("surveys").select("*").eq("building", TARGET_BUILDING).execute()
df = pd.DataFrame(response.data)
print(f"surveys: {len(df)} แถว")

survey_ids = df["id"].tolist()
hops_resp = client.table("traceroute_hops").select("*").in_("survey_id", survey_ids).execute()
df_hops = pd.DataFrame(hops_resp.data)
print(f"traceroute_hops: {len(df_hops)} แถว")

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def safe_float(val):
    try:
        return float(val) if val is not None else None
    except:
        return None

def safe_val(val, decimals=None):
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return "-"
    if decimals is not None:
        try:
            return str(round(float(val), decimals))
        except:
            return str(val)
    return str(val)

def bold_fill_cell(cell, text, bg="2C3E50", fg="FFFFFF", size=9):
    cell.value = text
    cell.font = Font(bold=True, color=fg, size=size, name="TH Sarabun New")
    cell.fill = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def thin_border():
    s = Side(style="thin", color="CCCCCC")
    return Border(left=s, right=s, top=s, bottom=s)

def add_word_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ {text} ]")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.name = "TH Sarabun New"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_cell(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def style_word_header(cell, text, bg="1B2A3B", fg="FFFFFF"):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.italic = False
    run.font.underline = False
    run.font.color.rgb = RGBColor.from_string(fg)
    run.font.size = Pt(10)
    run.font.name = "TH Sarabun New"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, bg)

def make_bar_chart(title, y_title, x_title, y_min=None, y_max=None, major_unit=None, width=26, height=14):
    c = BarChart()
    c.type = "col"
    c.grouping = "clustered"
    c.title = title
    c.style = 2
    c.width = width
    c.height = height
    c.y_axis.title = y_title
    c.y_axis.numFmt = "0"
    c.y_axis.delete = False
    c.y_axis.majorGridlines = None
    c.y_axis.tickLblPos = "low"
    if y_min is not None:
        c.y_axis.scaling.min = y_min
    if y_max is not None:
        c.y_axis.scaling.max = y_max
    if major_unit is not None:
        c.y_axis.majorUnit = major_unit
    c.x_axis.title = x_title
    c.x_axis.tickLblPos = "low"
    c.x_axis.noMultiLvlLbl = True
    return c

def make_line_chart(title, y_title, x_title, y_min=0, width=26, height=14):
    c = LineChart()
    c.title = title
    c.style = 2
    c.width = width
    c.height = height
    c.y_axis.title = y_title
    c.y_axis.numFmt = "0.0"
    c.y_axis.delete = False
    c.y_axis.majorGridlines = None
    c.y_axis.tickLblPos = "low"
    if y_min is not None:
        c.y_axis.scaling.min = y_min
    c.x_axis.title = x_title
    c.x_axis.tickLblPos = "low"
    c.x_axis.noMultiLvlLbl = True
    return c

# สร้างโฟลเดอร์ตามชื่อตึก
output_dir = os.path.join("Reports", TARGET_BUILDING)
os.makedirs(output_dir, exist_ok=True)

# ============================================================
# EXCEL
# ============================================================
wb = openpyxl.Workbook()
HEADER_FILL = PatternFill("solid", fgColor="2C3E50")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=9, name="TH Sarabun New")
DATA_FONT   = Font(size=9, name="TH Sarabun New")
rating_fills = {
    "Good":      PatternFill("solid", fgColor="D5F5E3"),
    "Excellent": PatternFill("solid", fgColor="A9DFBF"),
    "Poor":      PatternFill("solid", fgColor="FADBD8"),
}

# ----------------------------------------------------------
# Sheet 1: Data
# ----------------------------------------------------------
ws_data = wb.active
ws_data.title = "Data"

all_cols = [
    "room_point","note","floor","ssid","bssid","ap_vendor","band","radio_type",
    "channel","signal_percent","rssi_dbm","receive_rate_mbps","transmit_rate_mbps",
    "gateway_ip","ping_gateway_ms","ping_gateway_loss_pct","server_ip","trace_target",
    "ping_server_ms","ping_server_loss_pct","tcp_upload_mbps","tcp_download_mbps",
    "udp_target_bandwidth","udp_actual_mbps","udp_jitter_ms","udp_packetloss_pct",
    "co_channel_ap_count","adjacent_ap_count","strongest_neighbor_rssi",
    "noise_floor_dbm","snr_db","snr_quality","rating"
]
all_cols_th = [
    "ห้อง","จุด","ชั้น","SSID","BSSID","AP Vendor","Band","Radio Type",
    "Channel","Signal %","RSSI (dBm)","Receive Rate (Mbps)","Transmit Rate (Mbps)",
    "Gateway IP","Ping Gateway (ms)","Gateway Loss %","Server IP","Trace Target",
    "Ping Server (ms)","Server Loss %","TCP Upload (Mbps)","TCP Download (Mbps)",
    "UDP Target BW","UDP Actual (Mbps)","UDP Jitter (ms)","UDP Loss %",
    "Co-Channel AP","Adjacent AP","Neighbor RSSI",
    "Noise Floor (dBm)","SNR (dB)","SNR Quality","Rating"
]

for ci, h in enumerate(all_cols_th, 1):
    c = ws_data.cell(row=1, column=ci, value=h)
    c.fill = HEADER_FILL
    c.font = HEADER_FONT
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    c.border = thin_border()

for ri, (_, row) in enumerate(df.iterrows(), 2):
    rating = str(row.get("rating", ""))
    rfill = rating_fills.get(rating, PatternFill("solid", fgColor="FFFFFF"))
    for ci, col in enumerate(all_cols, 1):
        val = row.get(col)
        try:
            val = float(val) if val is not None and str(val).replace('.','').replace('-','').isdigit() else val
        except:
            pass
        c = ws_data.cell(row=ri, column=ci, value=val)
        c.fill = rfill
        c.font = DATA_FONT
        c.border = thin_border()
        c.alignment = Alignment(vertical="center")

ws_data.row_dimensions[1].height = 35
ws_data.freeze_panes = "A2"
for col in ws_data.columns:
    max_w = max((len(str(c.value or "")) for c in col), default=8) + 3
    ws_data.column_dimensions[get_column_letter(col[0].column)].width = min(max_w, 22)

# ----------------------------------------------------------
# Sheet 2: กราฟ RSSI
# ----------------------------------------------------------
ws_rssi = wb.create_sheet("กราฟ RSSI")

headers_rssi = ["ห้อง / จุด", "RSSI (dBm)"]
for ci, h in enumerate(headers_rssi, 1):
    c = ws_rssi.cell(1, ci, h)
    c.font = HEADER_FONT
    c.fill = HEADER_FILL
    c.alignment = Alignment(horizontal="center")

for i, (_, row) in enumerate(df.iterrows(), 2):
    ws_rssi.cell(i, 1, f"{row['room_point']} - {row.get('note','')}")
    ws_rssi.cell(i, 2, safe_float(row.get("rssi_dbm")))

ws_rssi.column_dimensions["A"].width = 30
ws_rssi.column_dimensions["B"].width = 14

rssi_vals = df["rssi_dbm"].dropna().astype(float)
rssi_min = int(rssi_vals.min()) - 5 if len(rssi_vals) else -80
rssi_max = 0

# RSSI bar chart
c1 = make_bar_chart(
    title="RSSI (dBm) ตามจุดสำรวจ",
    y_title="dBm",
    x_title="ห้อง / จุดสำรวจ",
    y_min=rssi_min,
    y_max=rssi_max,
    major_unit=5,
)
c1.add_data(Reference(ws_rssi, min_col=2, min_row=1, max_row=len(df)+1), titles_from_data=True)
c1.set_categories(Reference(ws_rssi, min_col=1, min_row=2, max_row=len(df)+1))
ws_rssi.add_chart(c1, "E2")



# ----------------------------------------------------------
# Sheet 3: กราฟ Speed
# ----------------------------------------------------------
ws_speed = wb.create_sheet("กราฟ Speed")

headers_spd = ["ห้อง", "TCP Upload (Mbps)", "TCP Download (Mbps)", "UDP Actual (Mbps)"]
for ci, h in enumerate(headers_spd, 1):
    c = ws_speed.cell(1, ci, h)
    c.font = HEADER_FONT
    c.fill = HEADER_FILL
    c.alignment = Alignment(horizontal="center")

for i, (_, row) in enumerate(df.iterrows(), 2):
    ws_speed.cell(i, 1, f"{row['room_point']} - {row.get('note','')}")
    ws_speed.cell(i, 2, safe_float(row.get("tcp_upload_mbps")))
    ws_speed.cell(i, 3, safe_float(row.get("tcp_download_mbps")))
    ws_speed.cell(i, 4, safe_float(row.get("udp_actual_mbps")))

ws_speed.column_dimensions["A"].width = 30

c3 = make_bar_chart(
    title="ความเร็ว TCP Upload / Download / UDP (Mbps)",
    y_title="Mbps",
    x_title="ห้อง / จุดสำรวจ",
    y_min=0,
    width=28,
    height=15,
)
c3.add_data(Reference(ws_speed, min_col=2, min_row=1, max_col=4, max_row=len(df)+1), titles_from_data=True)
c3.set_categories(Reference(ws_speed, min_col=1, min_row=2, max_row=len(df)+1))
ws_speed.add_chart(c3, "F2")

# ----------------------------------------------------------
# Sheet 4: กราฟ Latency
# ----------------------------------------------------------
ws_lat = wb.create_sheet("กราฟ Latency")

headers_lat = ["ห้อง", "Ping Gateway (ms)", "Ping Server (ms)", "Gateway Loss %", "Server Loss %"]
for ci, h in enumerate(headers_lat, 1):
    c = ws_lat.cell(1, ci, h)
    c.font = HEADER_FONT
    c.fill = HEADER_FILL
    c.alignment = Alignment(horizontal="center")

for i, (_, row) in enumerate(df.iterrows(), 2):
    ws_lat.cell(i, 1, f"{row['room_point']} - {row.get('note','')}")
    ws_lat.cell(i, 2, safe_float(row.get("ping_gateway_ms")))
    ws_lat.cell(i, 3, safe_float(row.get("ping_server_ms")))
    ws_lat.cell(i, 4, safe_float(row.get("ping_gateway_loss_pct")))
    ws_lat.cell(i, 5, safe_float(row.get("ping_server_loss_pct")))

ws_lat.column_dimensions["A"].width = 30

# คำนวณ y_max latency แบบตัด outlier (percentile 95)
lat_all = (
    df["ping_gateway_ms"].dropna().astype(float).tolist() +
    df["ping_server_ms"].dropna().astype(float).tolist()
)
lat_max = int(np.percentile(lat_all, 95)) + 20 if lat_all else 100

c4 = make_bar_chart(
    title="Latency — Ping Gateway vs Server (ms)",
    y_title="ms",
    x_title="ห้อง / จุดสำรวจ",
    y_min=0,
    y_max=lat_max,
    width=28,
    height=14,
)
c4.add_data(Reference(ws_lat, min_col=2, min_row=1, max_col=3, max_row=len(df)+1), titles_from_data=True)
c4.set_categories(Reference(ws_lat, min_col=1, min_row=2, max_row=len(df)+1))
ws_lat.add_chart(c4, "G2")

c5 = make_line_chart(
    title="Packet Loss % — Gateway & Server",
    y_title="%",
    x_title="ห้อง / จุดสำรวจ",
    y_min=0,
    width=28,
    height=14,
)
c5.add_data(Reference(ws_lat, min_col=4, min_row=1, max_col=5, max_row=len(df)+1), titles_from_data=True)
c5.set_categories(Reference(ws_lat, min_col=1, min_row=2, max_row=len(df)+1))
ws_lat.add_chart(c5, "G22")

# ----------------------------------------------------------
# Sheet 5: WinMTR Data
# ----------------------------------------------------------
if not df_hops.empty:
    ws_mtr = wb.create_sheet("WinMTR Data")
    mtr_cols = ["survey_id","room_point","hop_no","ip","loss_pct","sent","recv",
                "best_ms","avg_ms","worst_ms","last_ms"]
    mtr_th   = ["Survey ID","ห้อง","Hop #","IP","Loss %","Sent","Recv",
                "Best (ms)","Avg (ms)","Worst (ms)","Last (ms)"]
    for ci, h in enumerate(mtr_th, 1):
        c = ws_mtr.cell(1, ci, h)
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = Alignment(horizontal="center")
        c.border = thin_border()
    for ri, (_, row) in enumerate(df_hops.iterrows(), 2):
        for ci, col in enumerate(mtr_cols, 1):
            c = ws_mtr.cell(ri, ci, row.get(col))
            c.font = DATA_FONT
            c.border = thin_border()
    for col in ws_mtr.columns:
        max_w = max((len(str(c.value or "")) for c in col), default=8) + 3
        ws_mtr.column_dimensions[get_column_letter(col[0].column)].width = min(max_w, 20)
    ws_mtr.freeze_panes = "A2"



xlsx_name = os.path.join(output_dir, f"wifi_report_{TARGET_BUILDING}.xlsx")
wb.save(xlsx_name)
print(f"✅ Excel: {xlsx_name}")

# ============================================================
# WORD
# ============================================================
doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ----------------------------------------------------------
# หน้าปก
# ----------------------------------------------------------
# Heading 1: ชื่อคณะ (แก้ได้ภายหลัง)
p = doc.add_heading(f"ผลการสำรวจคุณภาพ WiFi", level=1)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

survey_date = pd.to_datetime(df["survey_timestamp"].iloc[0]).strftime("%d/%m/%Y") if "survey_timestamp" in df.columns else "-"
info = doc.add_paragraph(f"วันที่สำรวจ: {survey_date}  |  จำนวนจุดสำรวจ: {len(df)} จุด")
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

# Heading 2: อาคาร
doc.add_heading(f"อาคาร {TARGET_BUILDING}", level=2)

# ----------------------------------------------------------
# Heading 3: สรุปผลภาพรวม  (อยู่ใต้ อาคาร)
# ----------------------------------------------------------
total    = len(df)
passed   = len(df[df["rating"].isin(["Good","Excellent"])])
pass_pct = round(passed/total*100, 1) if total > 0 else 0
avg_rssi = round(df["rssi_dbm"].astype(float).mean(), 1)
problems = len(df[df["rating"] == "Poor"])

doc.add_heading("สรุปผลภาพรวม", level=3)
ts = doc.add_table(rows=2, cols=4)
ts.style = "Table Grid"
for i, (h_, v_) in enumerate(zip(
    ["จุดสำรวจทั้งหมด", "ผ่านเกณฑ์ (Good+)", "RSSI เฉลี่ย", "พบปัญหา (Poor)"],
    [str(total), f"{pass_pct}%", f"{avg_rssi} dBm", str(problems)]
)):
    style_word_header(ts.rows[0].cells[i], h_)
    ts.rows[1].cells[i].text = v_
    ts.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ----------------------------------------------------------
# Heading 3: กราฟสรุปภาพรวม  (อยู่ใต้ อาคาร)
# ----------------------------------------------------------
doc.add_heading("กราฟสรุปภาพรวม", level=3)

chart_placeholders = [
    ("ความแรงสัญญาณ RSSI (dBm) ตามจุดสำรวจ",
     f"Copy กราฟจาก {xlsx_name} → Sheet 'กราฟ RSSI' (กราฟบน)"),
    ("ความเร็ว TCP Upload / Download / UDP (Mbps)",
     f"Copy กราฟจาก {xlsx_name} → Sheet 'กราฟ Speed'"),
    ("Latency — Ping Gateway vs Server (ms)",
     f"Copy กราฟจาก {xlsx_name} → Sheet 'กราฟ Latency' (กราฟบน)"),
    ("Packet Loss % — Gateway & Server",
     f"Copy กราฟจาก {xlsx_name} → Sheet 'กราฟ Latency' (กราฟล่าง)"),
]
for chart_title, ph in chart_placeholders:
    doc.add_heading(chart_title, level=4)
    add_word_placeholder(doc, ph)
    doc.add_paragraph("")

doc.add_page_break()

# ----------------------------------------------------------
# field_map สำหรับตารางรายจุด
# ----------------------------------------------------------
field_map = [
    ("SSID",                    "ssid"),
    ("BSSID",                   "bssid"),
    ("AP Vendor",               "ap_vendor"),
    ("Band",                    "band"),
    ("Radio Type",              "radio_type"),
    ("Channel",                 "channel"),
    ("Signal %",                "signal_percent"),
    ("RSSI (dBm)",              "rssi_dbm"),
    ("Receive Rate (Mbps)",     "receive_rate_mbps"),
    ("Transmit Rate (Mbps)",    "transmit_rate_mbps"),
    ("Gateway IP",              "gateway_ip"),
    ("Ping Gateway (ms)",       "ping_gateway_ms"),
    ("Gateway Loss %",          "ping_gateway_loss_pct"),
    ("Server IP",               "server_ip"),
    ("Trace Target",            "trace_target"),
    ("Ping Server (ms)",        "ping_server_ms"),
    ("Server Loss %",           "ping_server_loss_pct"),
    ("TCP Upload (Mbps)",       "tcp_upload_mbps"),
    ("TCP Download (Mbps)",     "tcp_download_mbps"),
    ("UDP Target BW (Mbps)",    "udp_target_bandwidth"),
    ("UDP Actual (Mbps)",       "udp_actual_mbps"),
    ("UDP Jitter (ms)",         "udp_jitter_ms"),
    ("UDP Packet Loss %",       "udp_packetloss_pct"),
    ("Co-Channel AP",           "co_channel_ap_count"),
    ("Adjacent AP",             "adjacent_ap_count"),
    ("Strongest Neighbor RSSI", "strongest_neighbor_rssi"),
    ("Rating",                  "rating"),
]

# คอลัมน์ที่เว้นว่างไว้กรอกเอง
MANUAL_BLANK_COLS = {"co_channel_ap_count", "strongest_neighbor_rssi", "adjacent_ap_count"}

# ----------------------------------------------------------
# Heading 3: ชั้น → Heading 4: ห้อง  (อยู่ใต้ อาคาร)
# ----------------------------------------------------------
floors = df["floor"].fillna("ไม่ระบุชั้น").unique()
floors_sorted = sorted(floors, key=lambda x: (str(x) == "ไม่ระบุชั้น", x))

for floor_val in floors_sorted:
    df_floor = df[df["floor"].fillna("ไม่ระบุชั้น") == floor_val].reset_index(drop=True)
    doc.add_heading(f"ชั้น {floor_val}", level=3)

    rooms = df_floor["room_point"].unique()
    for room in rooms:
        df_room = df_floor[df_floor["room_point"] == room].reset_index(drop=True)
        n_pts   = len(df_room)

        doc.add_heading(f"ห้อง: {room}", level=4)

        # Heading 5: ข้อมูลทั่วไป
        doc.add_heading("ข้อมูลทั่วไปของห้อง", level=5)
        t_info = doc.add_table(rows=2, cols=3)
        t_info.style = "Table Grid"
        for ci, h_ in enumerate(["ห้อง", "ขนาดพื้นที่", "จำนวน AP"]):
            style_word_header(t_info.rows[0].cells[ci], h_)
        t_info.rows[1].cells[0].text = str(room)
        t_info.rows[1].cells[1].text = ""
        t_info.rows[1].cells[2].text = ""
        p_cap = doc.add_paragraph()
        run_cap = p_cap.add_run("อัตราส่วนจำนวน Client ที่รองรับได้ต่อ Access Point (Clients per AP):  ")
        run_cap.font.bold = True
        run_cap.font.size = Pt(10)
        run_cap.font.name = "TH Sarabun New"
        run_blank = p_cap.add_run("_______________")
        run_blank.font.size = Pt(10)
        run_blank.font.name = "TH Sarabun New"
        doc.add_paragraph("")

        # Heading 5: ข้อมูลการวัด WiFi -- แยกตาม SSID
        ssids_in_room = df_room["ssid"].unique().tolist()
        for ssid_val in ssids_in_room:
            df_ssid = df_room[df_room["ssid"] == ssid_val].reset_index(drop=True)
            n_ssid_pts = len(df_ssid)

            doc.add_heading(f"ข้อมูลการวัด WiFi (SSID: {ssid_val})", level=5)
            t_wifi = doc.add_table(rows=len(field_map)+2, cols=n_ssid_pts+1)
            t_wifi.style = "Table Grid"

            # แถวที่ 0 col0 + แถว 1 col0: merge แนวตั้ง → "รายการ"
            t_wifi.rows[0].cells[0].merge(t_wifi.rows[1].cells[0])
            hdr_cell = t_wifi.rows[0].cells[0]
            style_word_header(hdr_cell, "รายการ")
            # แถวที่ 0 col1..n: merge แนวนอน → "ตำแหน่งที่วัด"
            if n_ssid_pts > 1:
                t_wifi.rows[0].cells[1].merge(t_wifi.rows[0].cells[n_ssid_pts])
            style_word_header(t_wifi.rows[0].cells[1], "ตำแหน่งที่วัด")

            # แถวที่ 1: col0 ถูก merge ไปแล้ว → ใส่แค่ชื่อจุดแต่ละคอลัมน์
            for pt_i, (_, pt_row) in enumerate(df_ssid.iterrows()):
                style_word_header(
                    t_wifi.rows[1].cells[pt_i+1],
                    f"{pt_row.get('note', pt_i+1)}",
                    bg="2E4057",
                )

            for fi, (label, col) in enumerate(field_map, 2):
                cell0 = t_wifi.rows[fi].cells[0]
                # ถ้า cell0 ยัง reference ไปที่ merged header cell ให้ข้ามไป
                if cell0._tc is not hdr_cell._tc:
                    cell0.paragraphs[0].clear()
                    run0 = cell0.paragraphs[0].add_run(label)
                    run0.font.bold = True
                    run0.font.italic = False
                    run0.font.size = Pt(10)
                    run0.font.name = "TH Sarabun New"
                    cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for pt_i, (_, pt_row) in enumerate(df_ssid.iterrows()):
                    cell = t_wifi.rows[fi].cells[pt_i+1]
                    if col in MANUAL_BLANK_COLS:
                        cell.text = ""
                    else:
                        val = safe_val(pt_row.get(col))
                        cell.text = val
                        if col == "rating":
                            color_map = {"Good":"D5F5E3","Excellent":"A9DFBF","Poor":"FADBD8"}
                            shade_cell(cell, color_map.get(val, "FFFFFF"))
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.italic = False
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        cell.paragraphs[0].runs[0].font.name = "TH Sarabun New"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Heading 5: nperf (แยกตาม SSID)
            doc.add_heading("ผลการทดสอบ nperf", level=5)
            t_np = doc.add_table(rows=2, cols=4)
            t_np.style = "Table Grid"
            for ci, h_ in enumerate(["Download (Mbps)", "Upload (Mbps)", "Latency (ms)", "Jitter (ms)"]):
                style_word_header(t_np.rows[0].cells[ci], h_)
            for ci in range(4):
                t_np.rows[1].cells[ci].text = ""
            doc.add_paragraph("")

            # Heading 5: WinMTR Traceroute (แยกตาม SSID)
            doc.add_heading("ผล Traceroute", level=5)
            if not df_hops.empty:
                ssid_ids   = df_ssid["id"].tolist()
                df_ss_hops = df_hops[df_hops["survey_id"].isin(ssid_ids)].sort_values(["survey_id","hop_no"])
                if not df_ss_hops.empty:
                    mtr_f  = ["hop_no","ip","loss_pct","sent","recv","best_ms","avg_ms","worst_ms","last_ms"]
                    mtr_th = ["Hop #","IP","Loss %","Sent","Recv","Best (ms)","Avg (ms)","Worst (ms)","Last (ms)"]
                    t_mtr  = doc.add_table(rows=1, cols=len(mtr_f))
                    t_mtr.style = "Table Grid"
                    for ci, h_ in enumerate(mtr_th):
                        style_word_header(t_mtr.rows[0].cells[ci], h_)
                    for _, hop in df_ss_hops.iterrows():
                        r = t_mtr.add_row()
                        for ci, col in enumerate(mtr_f):
                            r.cells[ci].text = safe_val(hop.get(col))
                            if r.cells[ci].paragraphs[0].runs:
                                r.cells[ci].paragraphs[0].runs[0].font.italic = False
                                r.cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
                                r.cells[ci].paragraphs[0].runs[0].font.name = "TH Sarabun New"
                else:
                    doc.add_paragraph("ไม่พบข้อมูล WinMTR สำหรับ SSID นี้")
            else:
                doc.add_paragraph("ไม่มีข้อมูล WinMTR")
            doc.add_paragraph("")

        # Heading 5: Spectrum
        doc.add_heading("รูป Spectrum Analyzer", level=5)
        add_word_placeholder(doc, f"แทรกรูป Spectrum ห้อง {room} ที่นี่")
        doc.add_paragraph("")

        # Heading 5: รูปห้อง
        doc.add_heading("รูปภาพห้อง", level=5)
        add_word_placeholder(doc, f"แทรกรูปภาพห้อง {room} ที่นี่")

        doc.add_page_break()

# ----------------------------------------------------------
# Heading 3: ปัญหาที่ตรวจพบ  (อยู่ใต้ อาคาร)
# ----------------------------------------------------------
doc.add_heading("ปัญหาที่ตรวจพบ", level=3)
doc.add_paragraph("[กรอกรายละเอียดปัญหาที่ตรวจพบที่นี่]")

docx_name = os.path.join(output_dir, f"wifi_report_{TARGET_BUILDING}.docx")
doc.save(docx_name)
print(f"✅ Word:  {docx_name}")
print("✅ เสร็จสมบูรณ์!")