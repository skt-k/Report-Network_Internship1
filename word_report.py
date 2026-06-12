import os
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from utils import (
    FIELD_MAP,
    MANUAL_BLANK_COLS,
    safe_val,
    shade_cell,
    style_word_header,
)


def _apply_paragraph_style(paragraph, text):
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "TH Sarabun New"
    return run


def _add_placeholder(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"[ {text} ]")
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.name = "TH Sarabun New"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _create_table_row(table, cells):
    row = table.add_row()
    for idx, value in enumerate(cells):
        cell = row.cells[idx]
        cell.text = str(value)
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.name = "TH Sarabun New"
    return row


def create_word_report(df, df_hops, target_building, output_dir, excel_path=None):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    doc.add_heading("ผลการสำรวจคุณภาพ WiFi", level=1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    survey_date = "-"
    if "survey_timestamp" in df.columns and not df["survey_timestamp"].empty:
        try:
            survey_date = pd.to_datetime(df["survey_timestamp"].iloc[0]).strftime("%d/%m/%Y")
        except Exception:
            survey_date = "-"
    doc.add_paragraph(f"วันที่สำรวจ: {survey_date}  |  จำนวนจุดสำรวจ: {len(df)} จุด").alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_page_break()

    doc.add_heading(f"อาคาร {target_building}", level=2)
    total = len(df)
    passed = len(df[df["rating"].isin(["Good", "Excellent"])]) if "rating" in df.columns else 0
    pass_pct = round(passed / total * 100, 1) if total else 0
    avg_rssi = round(df["rssi_dbm"].astype(float).mean(), 1) if "rssi_dbm" in df.columns and not df["rssi_dbm"].dropna().empty else 0
    problems = len(df[df["rating"] == "Poor"]) if "rating" in df.columns else 0

    doc.add_heading("สรุปผลภาพรวม", level=3)
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = "Table Grid"
    for idx, header in enumerate(["จุดสำรวจทั้งหมด", "ผ่านเกณฑ์ (Good+)", "RSSI เฉลี่ย", "พบปัญหา (Poor)"]):
        style_word_header(summary_table.rows[0].cells[idx], header)
    summary_table.rows[1].cells[0].text = str(total)
    summary_table.rows[1].cells[1].text = f"{pass_pct}%"
    summary_table.rows[1].cells[2].text = f"{avg_rssi} dBm"
    summary_table.rows[1].cells[3].text = str(problems)
    for cell in summary_table.rows[1].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_heading("กราฟสรุปภาพรวม", level=3)
    placeholder_texts = [
        (
            "ความแรงสัญญาณ RSSI (dBm) ตามจุดสำรวจ",
            f"Copy กราฟจาก {excel_path} -> Sheet 'กราฟ RSSI' (กราฟบน)",
        ),
        (
            "ความเร็ว TCP Upload / Download / UDP (Mbps)",
            f"Copy กราฟจาก {excel_path} -> Sheet 'กราฟ Speed'",
        ),
        (
            "Latency — Ping Gateway vs Server (ms)",
            f"Copy กราฟจาก {excel_path} -> Sheet 'กราฟ Latency' (กราฟบน)",
        ),
        (
            "Packet Loss % — Gateway & Server",
            f"Copy กราฟจาก {excel_path} -> Sheet 'กราฟ Latency' (กราฟล่าง)",
        ),
    ]
    for title, text in placeholder_texts:
        doc.add_heading(title, level=4)
        _add_placeholder(doc, text)
        doc.add_paragraph("")
    doc.add_page_break()

    floors = df["floor"].fillna("ไม่ระบุชั้น").unique() if "floor" in df.columns else ["ไม่ระบุชั้น"]
    floors_sorted = sorted(floors, key=lambda value: (str(value) == "ไม่ระบุชั้น", value))

    for floor in floors_sorted:
        floor_filter = df["floor"].fillna("ไม่ระบุชั้น") == floor if "floor" in df.columns else pd.Series([True] * len(df))
        df_floor = df[floor_filter].reset_index(drop=True)
        doc.add_heading(f"ชั้น {floor}", level=3)
        rooms = df_floor["room_point"].unique() if "room_point" in df_floor.columns else ["ไม่ระบุห้อง"]
        for room in rooms:
            df_room = df_floor[df_floor["room_point"] == room].reset_index(drop=True)
            doc.add_heading(f"ห้อง: {room}", level=4)

            doc.add_heading("ข้อมูลทั่วไปของห้อง", level=5)
            t_info = doc.add_table(rows=2, cols=3)
            t_info.style = "Table Grid"
            for ci, header in enumerate(["ห้อง", "ขนาดพื้นที่", "จำนวน AP"]):
                style_word_header(t_info.rows[0].cells[ci], header)
            t_info.rows[1].cells[0].text = str(room)
            t_info.rows[1].cells[1].text = ""
            t_info.rows[1].cells[2].text = ""
            p_cap = doc.add_paragraph()
            _apply_paragraph_style(p_cap, "อัตราส่วนจำนวน Client ที่รองรับได้ต่อ Access Point (Clients per AP):  ")
            run_blank = p_cap.add_run("_______________")
            run_blank.font.size = Pt(10)
            run_blank.font.name = "TH Sarabun New"
            doc.add_paragraph("")

            ssids = df_room["ssid"].unique().tolist() if "ssid" in df_room.columns else ["ไม่ระบุ SSID"]
            for ssid in ssids:
                df_ssid = df_room[df_room["ssid"] == ssid].reset_index(drop=True) if "ssid" in df_room.columns else df_room
                doc.add_heading(f"ข้อมูลการวัด WiFi (SSID: {ssid})", level=5)
                t_wifi = doc.add_table(rows=len(FIELD_MAP) + 2, cols=len(df_ssid) + 1)
                t_wifi.style = "Table Grid"
                t_wifi.rows[0].cells[0].merge(t_wifi.rows[1].cells[0])
                style_word_header(t_wifi.rows[0].cells[0], "รายการ")
                if len(df_ssid) > 1:
                    t_wifi.rows[0].cells[1].merge(t_wifi.rows[0].cells[len(df_ssid)])
                style_word_header(t_wifi.rows[0].cells[1], "ตำแหน่งที่วัด")

                for point_index, row in enumerate(df_ssid.itertuples(index=False)):
                    style_word_header(
                        t_wifi.rows[1].cells[point_index + 1],
                        str(getattr(row, 'note', point_index + 1)),
                        bg="2E4057",
                    )

                for row_index, (label, column_name) in enumerate(FIELD_MAP, start=2):
                    cell0 = t_wifi.rows[row_index].cells[0]
                    if cell0._tc is not t_wifi.rows[0].cells[0]._tc:
                        cell0.paragraphs[0].clear()
                        run = cell0.paragraphs[0].add_run(label)
                        run.font.bold = True
                        run.font.italic = False
                        run.font.size = Pt(10)
                        run.font.name = "TH Sarabun New"
                        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for point_index, row in enumerate(df_ssid.itertuples(index=False)):
                        cell = t_wifi.rows[row_index].cells[point_index + 1]
                        if column_name in MANUAL_BLANK_COLS:
                            cell.text = ""
                        else:
                            value = safe_val(getattr(row, column_name, None))
                            cell.text = value
                            if column_name == "rating":
                                color = {
                                    "Good": "D5F5E3",
                                    "Excellent": "A9DFBF",
                                    "Poor": "FADBD8",
                                }.get(value, "FFFFFF")
                                shade_cell(cell, color)
                        if cell.paragraphs[0].runs:
                            run = cell.paragraphs[0].runs[0]
                            run.font.italic = False
                            run.font.size = Pt(10)
                            run.font.name = "TH Sarabun New"
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                doc.add_heading("ผลการทดสอบ nperf", level=5)
                t_np = doc.add_table(rows=2, cols=4)
                t_np.style = "Table Grid"
                for ci, heading in enumerate(["Download (Mbps)", "Upload (Mbps)", "Latency (ms)", "Jitter (ms)"]):
                    style_word_header(t_np.rows[0].cells[ci], heading)
                for ci in range(4):
                    t_np.rows[1].cells[ci].text = ""
                doc.add_paragraph("")

                doc.add_heading("ผล Traceroute", level=5)
                if not df_hops.empty and "id" in df_ssid.columns:
                    ssid_ids = df_ssid["id"].tolist()
                    df_ss_hops = df_hops[df_hops["survey_id"].isin(ssid_ids)].sort_values(["survey_id", "hop_no"])
                    if not df_ss_hops.empty:
                        mtr_fields = [
                            "hop_no",
                            "ip",
                            "loss_pct",
                            "sent",
                            "recv",
                            "best_ms",
                            "avg_ms",
                            "worst_ms",
                            "last_ms",
                        ]
                        mtr_headers = [
                            "Hop #",
                            "IP",
                            "Loss %",
                            "Sent",
                            "Recv",
                            "Best (ms)",
                            "Avg (ms)",
                            "Worst (ms)",
                            "Last (ms)",
                        ]
                        t_mtr = doc.add_table(rows=1, cols=len(mtr_fields))
                        t_mtr.style = "Table Grid"
                        for ci, heading in enumerate(mtr_headers):
                            style_word_header(t_mtr.rows[0].cells[ci], heading)
                        for _, hop in df_ss_hops.iterrows():
                            row = t_mtr.add_row()
                            for ci, field in enumerate(mtr_fields):
                                cell = row.cells[ci]
                                cell.text = safe_val(hop.get(field))
                                if cell.paragraphs[0].runs:
                                    run = cell.paragraphs[0].runs[0]
                                    run.font.italic = False
                                    run.font.size = Pt(10)
                                    run.font.name = "TH Sarabun New"
                    else:
                        doc.add_paragraph("ไม่พบข้อมูล WinMTR สำหรับ SSID นี้")
                else:
                    doc.add_paragraph("ไม่มีข้อมูล WinMTR")
                doc.add_paragraph("")

            doc.add_heading("รูป Spectrum Analyzer", level=5)
            _add_placeholder(doc, f"แทรกรูป Spectrum ห้อง {room} ที่นี่")
            doc.add_paragraph("")

            doc.add_heading("รูปภาพห้อง", level=5)
            _add_placeholder(doc, f"แทรกรูปภาพห้อง {room} ที่นี่")
            doc.add_page_break()

    doc.add_heading("ปัญหาที่ตรวจพบ", level=3)
    doc.add_paragraph("[กรอกรายละเอียดปัญหาที่ตรวจพบที่นี่]")

    word_path = os.path.join(output_dir, f"wifi_report_{target_building}.docx")
    doc.save(word_path)
    return word_path
